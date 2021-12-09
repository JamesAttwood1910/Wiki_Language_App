from docx import Document
from bing_image_downloader import downloader
import os
from PIL import Image
from docx.shared import Inches
import nltk
import bs4 as bs
import urllib.request
import re
import heapq
import requests
from bs4 import BeautifulSoup


document = Document()

document.add_heading('Welcome to your English activity worksheet')

document.add_heading('Activity 1', 
	level = 3)

paragraph = document.add_paragraph('Talk about each of the images (What, where, who, how, why). Use your imagination.')

#################  - Generate your images

string = input("Search query: ")

number_photos = input('How many photos? ')

folder_name = input('Choose your folder name: ')

downloader.download(string, limit= int(number_photos),  
    output_dir=folder_name, adult_filter_off=True, 
    force_replace=False, timeout=60, verbose=True)

print("Photos saved in", folder_name,"/",string)

extensions = ('jpeg', 'png')

directory = "{}/{}".format(folder_name, string)

files_in_directory = os.listdir(directory)

filtered_files = [file for file in files_in_directory if file.endswith(extensions)]

for file in filtered_files:
    path_to_file = os.path.join(directory, file)
    os.remove(path_to_file)


files_in_directory_filtered = os.listdir(directory)

for file in files_in_directory_filtered:
	document.add_picture(directory + '/' + file, width = Inches(6.0))
	print(

##############################################

		)
#   img = Image.open(directory + '/' + file)
#   img.show()


######################################## - text

document.add_heading('Activity 2', 
	level = 3)

document.add_paragraph('Read the below article.')

wiki_article = input("Choose your wikipedia article:")

url = wiki_article
result = requests.get(url)
doc = BeautifulSoup(result.text, "html.parser")

titles = []
for x in doc.find_all('span', class_ = 'mw-headline'):
  if x.string == 'External links':
    continue
  if x.string == 'References':
    continue
  if x.string == 'See also':
    continue
  if x.string == 'Further reading':
    continue
  titles.append(x.string)

for title in titles:
  document.add_heading(title, level =4)
  
  paragraphs = title.find_all_next('p', limit = 2)
  for p in paragraphs: 
    text = str(p.get_text())
    text = re.sub(r'\[[0-9]*\]',' ',text)
    document.add_paragraph(text)
    print(

    	)




















document.save('word_doc_{}.docx'.format(string))







